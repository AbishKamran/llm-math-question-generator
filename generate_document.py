from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from question_generator import MathQuestionGenerator
from question_analytics import generate_analytics_report
from similarity_checker import QuestionSimilarityChecker
import json

def create_enhanced_word_document():
    """Create an enhanced Word document with analytics and quality checks"""
    generator = MathQuestionGenerator(difficulty_adaptive=True, latex_support=True)
    
    # Generate multiple questions for better analysis
    questions = []
    for i in range(4):  # Generate 4 questions
        if i % 2 == 0:
            q = generator.generate_counting_question()
        else:
            q = generator.generate_geometry_question()
        questions.append(q)
    
    # Run analytics and similarity checks
    analytics = generate_analytics_report(questions)
    similarity_checker = QuestionSimilarityChecker()
    similarity_report = similarity_checker.batch_similarity_check(questions)

    # Create enhanced Word document
    doc = Document()
    
    # Title page
    title = doc.add_heading('🧮 AI-Generated Mathematical Reasoning Assessment', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Advanced Assessment with Quality Analytics and Plagiarism Detection')
    doc.add_paragraph(f'Generated on: {analytics["generated_at"]}')
    doc.add_paragraph('\n')
    
    # Analytics summary
    doc.add_heading('📊 Quality Analytics Report', level=1)
    summary = analytics['summary']
    
    analytics_table = doc.add_table(rows=4, cols=2)
    analytics_table.style = 'Table Grid'
    
    metrics = [
        ('Quality Score', f"{summary['quality_score']*100:.1f}%"),
        ('Average Readability', f"{summary['avg_readability']:.1f}"),
        ('Engagement Score', f"{summary['avg_engagement']*100:.1f}%"),
        ('Total Questions', str(summary['total_questions']))
    ]
    
    for i, (metric, value) in enumerate(metrics):
        row = analytics_table.rows[i]
        row.cells[0].text = metric
        row.cells[1].text = value
    
    # Similarity report
    doc.add_heading('🔍 Plagiarism & Similarity Analysis', level=1)
    sim_summary = similarity_report['summary']
    
    doc.add_paragraph(f"Average Similarity: {sim_summary['average_similarity']*100:.1f}%")
    doc.add_paragraph(f"Maximum Similarity: {sim_summary['maximum_similarity']*100:.1f}%")
    doc.add_paragraph(f"High-Risk Pairs: {sim_summary['high_risk_pairs']}")
    
    for rec in similarity_report['recommendations']:
        doc.add_paragraph(f"• {rec}", style='List Bullet')
    
    doc.add_page_break()

    # Add all questions with enhanced formatting
    for i, q in enumerate(questions, 1):
        doc.add_heading(f'Question {i}', level=1)
        
        # Question text with difficulty indicator
        question_para = doc.add_paragraph()
        question_para.add_run(f"[{q['difficulty'].upper()}] ").bold = True
        question_para.add_run(q['question'])
        
        # Add table if present
        if 'table' in q:
            # Parse table data (simplified)
            doc.add_paragraph("\nReference Table:")
            table_lines = q['table'].split('\n')
            if len(table_lines) >= 3:  # Has header and data
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Add sample data (simplified for demo)
                if 'Sandwich' in q['table']:
                    headers = ['Sandwich', 'Drink']
                    data = [['Turkey', 'Water'], ['Ham', 'Juice'], ['Veggie', 'Milk'], ['Chicken', '']]
                else:
                    headers = ['Item 1', 'Item 2']
                    data = [['Option A', 'Choice 1'], ['Option B', 'Choice 2']]
                
                hdr_cells = table.rows[0].cells
                for j, header in enumerate(headers):
                    hdr_cells[j].text = header
                
                for row_data in data:
                    row_cells = table.add_row().cells
                    for j, cell_data in enumerate(row_data):
                        row_cells[j].text = cell_data
        
        doc.add_paragraph("\nOptions:")
        # Add options with correct answer highlighting
        for letter, option in q['options']:
            para = doc.add_paragraph()
            if letter == q['correct']:
                run = para.add_run(f"({letter}) {option} ✓")
                run.bold = True
                run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
            else:
                para.add_run(f"({letter}) {option}")
        
        # Enhanced explanation with formatting
        doc.add_paragraph("\nExplanation:", style='Heading 3')
        doc.add_paragraph(q['explanation'])
        
        # Curriculum info
        curriculum_para = doc.add_paragraph()
        curriculum_para.add_run("Curriculum: ").bold = True
        curriculum_para.add_run(f"{q['subject']} → {q['unit']} → {q['topic']}")
        
        # Quality metrics for this question
        if i <= len(analytics['individual_analyses']):
            analysis = analytics['individual_analyses'][i-1]
            metrics_para = doc.add_paragraph()
            metrics_para.add_run("Quality Metrics: ").bold = True
            metrics_para.add_run(f"Readability: {analysis['readability_score']:.1f}, ")
            metrics_para.add_run(f"Engagement: {analysis['engagement_score']*100:.1f}%")
        
        doc.add_paragraph("\n" + "="*50 + "\n")

    # Save enhanced document
    doc.save('Enhanced_Math_Questions.docx')
    print("Enhanced Word document created: Enhanced_Math_Questions.docx")
    
    # Save analytics report
    with open('analytics_report.json', 'w') as f:
        json.dump({
            'analytics': analytics,
            'similarity_report': similarity_report
        }, f, indent=2)
    print("Analytics report saved: analytics_report.json")
    
    # Also save formatted text version
    output = ""
    for i, q in enumerate(questions, 1):
        output += generator.format_question(q, i, "Enhanced Mathematical Reasoning Assessment" if i == 1 else "")
    
    with open('questions_formatted.txt', 'w', encoding='utf-8') as f:
        f.write(output)
    print("Formatted text file created: questions_formatted.txt")
    
    return questions, analytics, similarity_report

def create_word_document():
    """Backward compatibility function"""
    return create_enhanced_word_document()

if __name__ == "__main__":
    questions, analytics, similarity = create_enhanced_word_document()
    print(f"\n[ANALYTICS] Quality Score: {analytics['summary']['quality_score']*100:.1f}%")
    print(f"[SIMILARITY] Similarity Check: {similarity['summary']['high_risk_pairs']} high-risk pairs found")
    print("\n[SUCCESS] Enhanced document generation complete!")